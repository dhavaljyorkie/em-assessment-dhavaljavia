import io
import logging

from docx import Document

from src.parsers.base import BaseParser

logger = logging.getLogger(__name__)


class DocxParser(BaseParser):
    """
    Extracts text from .docx files using python-docx.
    Preserves paragraph order and includes table cell text.
    """

    def parse(self, data: bytes, filename: str = "") -> str:
        try:
            doc = Document(io.BytesIO(data))
            parts: list[str] = []

            # Body paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Table cells (resumes often use tables for layout)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            parts.append(cell.text.strip())

            return "\n".join(parts)
        except Exception as exc:
            logger.error("DocxParser: failed to parse '%s': %s", filename, exc)
            return ""
