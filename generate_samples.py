#!/usr/bin/env python3
"""
Generate synthetic sample resumes (DOCX) and a job description (TXT)
for Phase 7 integration testing.
Run: python generate_samples.py

Produces 20 diverse candidates:
  - Software Engineering: Python/Go, Java/.NET, iOS/Swift, Android/Kotlin,
    Embedded C/C++, Ruby on Rails, Rust/Systems, QA/SDET, Web3/Solidity
  - Engineering Management (strong and light)
  - ML / Data Science
  - DevOps / SRE
  - Security
  - Human Resources / Talent Acquisition
  - Finance / FP&A
  - Product Management
  - Marketing / Growth
  - Business Operations
"""

import os
from docx import Document

RESUMES_DIR = os.path.join(os.path.dirname(__file__), "data", "resumes")
JOBS_DIR = os.path.join(os.path.dirname(__file__), "data", "jobs")
os.makedirs(RESUMES_DIR, exist_ok=True)
os.makedirs(JOBS_DIR, exist_ok=True)

CANDIDATES = [
    # ── Engineering Management ────────────────────────────────────────────────
    {
        "name": "Maria Santos",
        "email": "maria.santos@example.com",
        "phone": "+55 11 99999 0808",
        "location": "São Paulo, Brazil (Remote)",
        "summary": (
            "Engineering Manager with 12 years in software engineering and 5 years in "
            "management. Scaled organisations from 5 to 50 engineers. Expert at hiring, "
            "roadmapping, and cross-functional delivery. Bilingual (English/Portuguese)."
        ),
        "skills": "Python, Java, AWS, Team Building, Roadmapping, Agile, OKRs, Hiring, Performance Reviews",
        "experience": [
            ("VP of Engineering", "Nubank (contractor)", "2023 – Present",
             "Led platform engineering group (30 engineers, 4 teams). Defined 2-year tech roadmap. Reduced time-to-production by 30%."),
            ("Engineering Manager", "iFood", "2020 – 2023",
             "Built Logistics Platform team from 4 to 18 engineers. Shipped real-time driver tracking feature. MTTR reduced 60%."),
            ("Senior Software Engineer", "Globo", "2015 – 2020",
             "Backend engineer on streaming platform. Migrated monolith to microservices using Python and AWS."),
        ],
        "education": [("B.Sc. Computer Science", "Universidade de São Paulo", "2012")],
    },
    {
        "name": "Ben Okafor",
        "email": "ben.okafor@example.com",
        "phone": "+44 7700 900123",
        "location": "London, UK",
        "summary": (
            "Engineering Manager with 10 years in software engineering and 3 years in "
            "people management. Led teams of 8–12 engineers across Payments and Core Platform. "
            "Strong at hiring, technical strategy, and delivery."
        ),
        "skills": "Python, TypeScript, AWS, Team Leadership, Agile, OKRs, Hiring, System Design",
        "experience": [
            ("Engineering Manager", "Monzo", "2022 – Present",
             "Managed Payments Platform team (10 engineers). Shipped Open Banking integration and reduced incident MTTR by 55%."),
            ("Senior Engineer → Tech Lead", "Revolut", "2018 – 2022",
             "Technical lead on FX trading feature. Grew team from 3 to 8 engineers. Introduced RFC-driven design process."),
            ("Software Engineer", "Thoughtworks", "2016 – 2018",
             "Consultant on financial services projects across UK and Germany."),
        ],
        "education": [("M.Eng. Software Engineering", "Imperial College London", "2015")],
    },

    # ── SWE: Python / Go / Backend ────────────────────────────────────────────
    {
        "name": "Alice Chen",
        "email": "alice.chen@example.com",
        "phone": "+1 415 555 0101",
        "location": "San Francisco, CA",
        "summary": (
            "Senior Software Engineer with 8 years of experience in backend systems, "
            "microservices, and distributed architectures. Focused on scalable systems and "
            "developer tooling."
        ),
        "skills": "Python, Go, Kubernetes, AWS, PostgreSQL, Redis, gRPC, Docker, Terraform",
        "experience": [
            ("Senior Software Engineer", "Stripe", "2021 – Present",
             "Led backend development of payment reconciliation service handling $2B/day. Reduced P99 latency by 40%."),
            ("Software Engineer", "Lyft", "2018 – 2021",
             "Built real-time ride-matching service in Go. Owned CI/CD pipeline migration to Kubernetes."),
        ],
        "education": [("B.Sc. Computer Science", "UC Berkeley", "2016")],
    },
    {
        "name": "Carlos Mendez",
        "email": "carlos.mendez@example.com",
        "phone": "+1 512 555 0303",
        "location": "Austin, TX",
        "summary": (
            "DevOps / Platform Engineer with 7 years building and operating large-scale "
            "cloud infrastructure. Expert in Kubernetes, Terraform, and reliability engineering."
        ),
        "skills": "Kubernetes, Terraform, AWS, GCP, Prometheus, Grafana, Helm, Python, Go, Bash",
        "experience": [
            ("Staff Platform Engineer", "HashiCorp", "2022 – Present",
             "Owned Kubernetes platform for 200+ microservices. Reduced deployment time from 45 min to 8 min."),
            ("Site Reliability Engineer", "GitHub", "2019 – 2022",
             "On-call for github.com. Drove 99.99% uptime SLA. Built self-healing infrastructure with Kubernetes operators."),
        ],
        "education": [("B.Sc. Computer Engineering", "UT Austin", "2017")],
    },

    # ── SWE: Java / JVM ──────────────────────────────────────────────────────
    {
        "name": "Rahul Sharma",
        "email": "rahul.sharma@example.com",
        "phone": "+91 98765 00909",
        "location": "Bangalore, India (Remote)",
        "summary": (
            "Staff Software Engineer with 11 years of experience in high-throughput distributed "
            "systems. Tech lead for 15 engineers. Holds 2 patents in distributed locking."
        ),
        "skills": "Java, Scala, Apache Kafka, Apache Spark, Cassandra, AWS, System Design, Technical Leadership",
        "experience": [
            ("Staff Software Engineer", "Flipkart", "2019 – Present",
             "Designed inventory management system handling 10M SKUs. Tech lead for 15 engineers. Reduced p99 checkout latency by 50%."),
            ("Senior Software Engineer", "InMobi", "2016 – 2019",
             "Built real-time bidding engine processing 500K requests/sec on Java + Kafka."),
            ("Software Engineer", "Infosys", "2014 – 2016",
             "Enterprise Java development (Spring Boot, Hibernate) for tier-1 banking clients."),
        ],
        "education": [("B.Tech. Computer Science", "IIT Bombay", "2013")],
    },
    {
        "name": "Aisha Kamara",
        "email": "aisha.kamara@example.com",
        "phone": "+1 347 555 0606",
        "location": "Remote (US East)",
        "summary": (
            "Backend Engineer specialising in event-driven architectures and data-intensive "
            "applications. 6 years with Kafka, Flink, and distributed databases."
        ),
        "skills": "Java, Kotlin, Apache Kafka, Apache Flink, Cassandra, PostgreSQL, AWS, Spring Boot",
        "experience": [
            ("Senior Backend Engineer", "Confluent", "2021 – Present",
             "Built Kafka Streams pipeline handling 5M events/sec. Open source contributor to Apache Flink."),
            ("Software Engineer", "DataStax", "2019 – 2021",
             "Cassandra performance optimisation. Led migration to multi-region active-active setup."),
        ],
        "education": [("B.Sc. Computer Science", "Howard University", "2018")],
    },

    # ── SWE: .NET / C# ───────────────────────────────────────────────────────
    {
        "name": "Oliver Braun",
        "email": "oliver.braun@example.com",
        "phone": "+49 176 555 0110",
        "location": "Berlin, Germany (Remote)",
        "summary": (
            "Senior .NET Engineer with 9 years of experience building enterprise SaaS products "
            "on C# and Azure. Expert in event sourcing, CQRS, DDD, and microservices."
        ),
        "skills": "C#, .NET 8, ASP.NET Core, Azure, Azure Service Bus, SQL Server, Entity Framework, Docker, CQRS, Event Sourcing",
        "experience": [
            ("Senior Software Engineer", "SAP", "2021 – Present",
             "Built order management microservices in .NET 8 / C#. Led migration from on-prem SQL Server to Azure SQL."),
            ("Software Engineer", "Zalando", "2018 – 2021",
             "Developed checkout and pricing services using CQRS and event sourcing patterns. Served 30M users."),
            ("Junior .NET Developer", "msg systems", "2016 – 2018",
             "Enterprise insurance software development in C# / WPF."),
        ],
        "education": [("B.Sc. Computer Science", "TU Berlin", "2015")],
    },

    # ── SWE: iOS / Swift ─────────────────────────────────────────────────────
    {
        "name": "Yuki Tanaka",
        "email": "yuki.tanaka@example.com",
        "phone": "+1 408 555 0220",
        "location": "San Jose, CA",
        "summary": (
            "iOS Engineer with 7 years building consumer apps for millions of users. "
            "Deep expertise in Swift, SwiftUI, and the Apple ecosystem. "
            "Experienced with Core Data, ARKit, and App Store optimisation."
        ),
        "skills": "Swift, SwiftUI, Objective-C, Xcode, Core Data, ARKit, XCTest, Fastlane, Firebase, REST APIs",
        "experience": [
            ("Senior iOS Engineer", "Airbnb", "2022 – Present",
             "Built SwiftUI search and discovery features for 70M MAU iOS app. Reduced app launch time by 25%."),
            ("iOS Developer", "Pinterest", "2019 – 2022",
             "Shipped Idea Pins (AR camera) feature. Led performance improvements that increased session length 18%."),
        ],
        "education": [("B.Sc. Computer Science", "Keio University", "2017")],
    },

    # ── SWE: Android / Kotlin ─────────────────────────────────────────────────
    {
        "name": "Kwame Asante",
        "email": "kwame.asante@example.com",
        "phone": "+1 646 555 0330",
        "location": "New York, NY",
        "summary": (
            "Android Engineer with 8 years of experience shipping high-quality apps on the "
            "Google Play Store. Expert in Kotlin, Jetpack Compose, and Android architecture patterns."
        ),
        "skills": "Kotlin, Java, Jetpack Compose, Android SDK, Room, Coroutines, Hilt, Retrofit, Firebase, Gradle",
        "experience": [
            ("Senior Android Engineer", "Spotify", "2021 – Present",
             "Rebuilt Spotify Android home feed in Jetpack Compose. Reduced ANR rate by 42%."),
            ("Android Developer", "Duolingo", "2018 – 2021",
             "Built gamification features for Android (streaks, leagues). 200M+ install base."),
        ],
        "education": [("B.Sc. Computer Science", "NYU Tandon", "2016")],
    },

    # ── SWE: Ruby on Rails ────────────────────────────────────────────────────
    {
        "name": "Chloe Dupont",
        "email": "chloe.dupont@example.com",
        "phone": "+33 6 55 01 04 04",
        "location": "Paris, France (Remote)",
        "summary": (
            "Full-stack Rails engineer with 6 years of experience at high-growth SaaS startups. "
            "Comfortable across the stack: Ruby backend, React/Hotwire frontend, PostgreSQL."
        ),
        "skills": "Ruby, Ruby on Rails, Hotwire, Turbo, React, PostgreSQL, Redis, Sidekiq, Docker, Heroku, AWS",
        "experience": [
            ("Senior Software Engineer", "Pennylane", "2022 – Present",
             "Built accounting automation features in Rails 7. Reduced month-end close time by 40% for 10K+ SMBs."),
            ("Software Engineer", "Alan", "2019 – 2022",
             "Developed health insurance member portal (Rails + React). Shipped claims management module."),
        ],
        "education": [("Engineering Degree (Software)", "École 42 Paris", "2018")],
    },

    # ── SWE: Rust / Systems ───────────────────────────────────────────────────
    {
        "name": "Elena Volkov",
        "email": "elena.volkov@example.com",
        "phone": "+7 916 555 0550",
        "location": "Amsterdam, Netherlands (Remote)",
        "summary": (
            "Systems Software Engineer with 8 years of experience, last 4 in Rust. "
            "Specialised in database internals, network protocols, and performance-critical "
            "infrastructure software."
        ),
        "skills": "Rust, C++, C, Linux, TCP/IP, gRPC, eBPF, LLVM, async-std, tokio, PostgreSQL internals",
        "experience": [
            ("Software Engineer", "CockroachDB", "2021 – Present",
             "Contributed to storage engine (Pebble) in Rust/Go. Improved write throughput by 25% through WAL optimisation."),
            ("Systems Engineer", "Cloudflare", "2018 – 2021",
             "Built eBPF-based DDoS mitigation in Rust. Processed 50M packets/sec per node."),
        ],
        "education": [("M.Sc. Computer Systems", "Moscow State University", "2016")],
    },

    # ── SWE: Embedded / C/C++ ─────────────────────────────────────────────────
    {
        "name": "Tariq Hassan",
        "email": "tariq.hassan@example.com",
        "phone": "+971 50 555 0660",
        "location": "Dubai, UAE (Remote)",
        "summary": (
            "Embedded Systems Engineer with 10 years of experience in firmware development for "
            "IoT devices, automotive ECUs, and real-time operating systems."
        ),
        "skills": "C, C++, RTOS (FreeRTOS, Zephyr), CAN bus, I2C, SPI, ARM Cortex-M, Python, Jenkins, JTAG debugging",
        "experience": [
            ("Senior Firmware Engineer", "Mobileye", "2020 – Present",
             "Developed ADAS sensor fusion firmware in C++ on ARM Cortex-A. Safety-critical (ISO 26262 ASIL-B)."),
            ("Embedded Software Engineer", "Bosch", "2016 – 2020",
             "Firmware for automotive brake ECUs. CAN/LIN protocol implementation in C."),
        ],
        "education": [("B.Sc. Electrical & Computer Engineering", "American University of Sharjah", "2014")],
    },

    # ── SWE: Full Stack / Frontend ────────────────────────────────────────────
    {
        "name": "Sophie Williams",
        "email": "sophie.williams@example.com",
        "phone": "+1 212 555 0404",
        "location": "New York, NY",
        "summary": (
            "Product-focussed Full Stack Engineer with 5 years experience at high-growth "
            "startups. Strong React/TypeScript frontend with Node.js/GraphQL backend skills."
        ),
        "skills": "TypeScript, React, Next.js, Node.js, GraphQL, PostgreSQL, Figma, Tailwind CSS, AWS",
        "experience": [
            ("Senior Full Stack Engineer", "Notion", "2022 – Present",
             "Core product team. Built real-time collaborative editing features. Reduced first-load time by 35%."),
            ("Full Stack Engineer", "Airtable", "2020 – 2022",
             "Built integrations platform (Zapier, Slack, GitHub). 200+ third-party integrations."),
        ],
        "education": [("B.A. Computer Science & Design", "Brown University", "2019")],
    },

    # ── SWE: QA / SDET ───────────────────────────────────────────────────────
    {
        "name": "Nadia Kovacs",
        "email": "nadia.kovacs@example.com",
        "phone": "+36 30 555 0770",
        "location": "Budapest, Hungary (Remote)",
        "summary": (
            "Senior QA Engineer / SDET with 7 years building test automation frameworks for "
            "web, mobile, and API layers. Advocate for quality-first engineering culture."
        ),
        "skills": "Python, Java, Selenium, Playwright, Appium, pytest, JUnit, Allure, Jenkins, k6, Postman, BDD/Gherkin",
        "experience": [
            ("Senior SDET", "LogMeIn", "2021 – Present",
             "Built Playwright + Python E2E framework from scratch. 80% reduction in manual regression time."),
            ("QA Engineer", "Prezi", "2018 – 2021",
             "API and UI test automation for presentation platform. Integrated tests into CI pipeline (Jenkins)."),
        ],
        "education": [("B.Sc. Software Engineering", "Budapest University of Technology", "2017")],
    },

    # ── SWE: ML / Data Science ────────────────────────────────────────────────
    {
        "name": "Priya Nair",
        "email": "priya.nair@example.com",
        "phone": "+1 650 555 0202",
        "location": "Seattle, WA",
        "summary": (
            "ML Engineer with 6 years building recommendation systems and NLP pipelines. "
            "Expertise in taking models from research to production at scale."
        ),
        "skills": "Python, PyTorch, TensorFlow, Hugging Face, Apache Spark, Kafka, AWS SageMaker, Airflow, SQL, MLflow",
        "experience": [
            ("ML Engineer II", "Amazon", "2021 – Present",
             "Built real-time product recommendation engine serving 50M daily users. 12% lift in CTR."),
            ("Data Scientist", "Zillow", "2019 – 2021",
             "Developed NLP model for automated property description generation. Deployed on SageMaker."),
        ],
        "education": [("M.Sc. Machine Learning", "Carnegie Mellon University", "2019")],
    },

    # ── SWE: Security ─────────────────────────────────────────────────────────
    {
        "name": "James Park",
        "email": "james.park@example.com",
        "phone": "+1 408 555 0505",
        "location": "San Jose, CA",
        "summary": (
            "Principal Security Engineer with 9 years in application security, threat "
            "modelling, and compliance. Deep expertise in OWASP Top 10, SOC2, and secure SDLC."
        ),
        "skills": "Python, Rust, AWS Security, IAM, Penetration Testing, SAST/DAST, OWASP, SOC2, ISO 27001, Burp Suite",
        "experience": [
            ("Principal Security Engineer", "Okta", "2020 – Present",
             "Led product security for Identity Cloud. Achieved SOC2 Type II and ISO 27001 certifications."),
            ("Application Security Engineer", "Palo Alto Networks", "2017 – 2020",
             "Threat modelling for next-gen firewalls. Embedded in 4 product teams."),
        ],
        "education": [("B.Sc. Computer Science (Security)", "Georgia Tech", "2015")],
    },

    # ── Web3 / Blockchain ─────────────────────────────────────────────────────
    {
        "name": "Diego Reyes",
        "email": "diego.reyes@example.com",
        "phone": "+34 611 555 0880",
        "location": "Madrid, Spain (Remote)",
        "summary": (
            "Blockchain / Web3 developer with 5 years building smart contracts and DeFi "
            "protocols. Experience with Ethereum, Solana, and Layer-2 scaling solutions."
        ),
        "skills": "Solidity, Rust, TypeScript, Hardhat, Foundry, Ethers.js, Solana (Anchor), IPFS, The Graph, AWS",
        "experience": [
            ("Smart Contract Engineer", "Uniswap Labs", "2022 – Present",
             "Built Uniswap v4 hook framework in Solidity. Audited by Trail of Bits. $2B TVL."),
            ("Web3 Developer", "Polygon", "2020 – 2022",
             "Developed Polygon Bridge contracts and SDK. Reduced bridge gas fees by 60%."),
        ],
        "education": [("B.Sc. Computer Science", "Universidad Politécnica de Madrid", "2019")],
    },

    # ── Human Resources ───────────────────────────────────────────────────────
    {
        "name": "Rachel Kim",
        "email": "rachel.kim@example.com",
        "phone": "+1 650 555 0990",
        "location": "San Francisco, CA",
        "summary": (
            "Senior HR Business Partner and Talent Acquisition leader with 9 years at "
            "high-growth tech companies. Expert in full-cycle recruiting, L&D, performance "
            "management, and DEI programmes."
        ),
        "skills": "Talent Acquisition, HRBP, Performance Management, L&D, DEI, Workday, Lever, Greenhouse, Compensation Benchmarking, Employment Law",
        "experience": [
            ("Senior HR Business Partner", "Databricks", "2021 – Present",
             "HRBP for Engineering org (300+ employees). Reduced regrettable attrition by 22%. Led bi-annual performance calibration."),
            ("Talent Acquisition Manager", "Lyft", "2018 – 2021",
             "Managed team of 8 technical recruiters. Hired 120+ engineers/year. Built university recruiting programme."),
            ("HR Generalist", "Box", "2016 – 2018",
             "Supported onboarding, employee relations, and HR operations for APAC region."),
        ],
        "education": [
            ("M.S. Human Resources Management", "Cornell ILR School", "2015"),
            ("B.A. Psychology", "UCLA", "2013"),
        ],
    },

    # ── Finance ───────────────────────────────────────────────────────────────
    {
        "name": "Marcus Webb",
        "email": "marcus.webb@example.com",
        "phone": "+1 212 555 1010",
        "location": "New York, NY",
        "summary": (
            "Finance Manager / FP&A Lead with 10 years at public tech companies. "
            "Expert in financial modelling, variance analysis, and partnering with "
            "Engineering and Product leaders on budget and headcount planning."
        ),
        "skills": "Financial Modelling, FP&A, Excel (advanced), SQL, Tableau, Anaplan, Workday Adaptive, GAAP, Investor Relations, Headcount Planning",
        "experience": [
            ("Finance Manager — R&D FP&A", "Salesforce", "2021 – Present",
             "Owned $800M R&D budget model. Partnered with CTO office on annual planning. Variance analysis and monthly forecast."),
            ("Senior Financial Analyst", "Twilio", "2018 – 2021",
             "Built 3-statement financial model for IPO readiness. Supported Series F–IPO investor relations process."),
            ("Financial Analyst", "JP Morgan", "2015 – 2018",
             "DCF valuation and M&A advisory for technology sector clients."),
        ],
        "education": [
            ("MBA (Finance)", "Wharton School, University of Pennsylvania", "2015"),
            ("B.Sc. Economics", "Duke University", "2013"),
        ],
    },

    # ── Product Management ────────────────────────────────────────────────────
    {
        "name": "Fatima Al-Rashid",
        "email": "fatima.alrashid@example.com",
        "phone": "+971 52 555 1120",
        "location": "Dubai, UAE (Remote)",
        "summary": (
            "Senior Product Manager with 8 years building developer tools and platform products. "
            "Adept at working with engineering, design, and data science to ship high-impact "
            "features for B2B SaaS companies."
        ),
        "skills": "Product Strategy, Roadmapping, User Research, A/B Testing, SQL, Figma, JIRA, Amplitude, Mixpanel, Go-To-Market",
        "experience": [
            ("Senior Product Manager", "GitHub", "2022 – Present",
             "PM for GitHub Copilot Enterprise. Grew adoption from 0 to 50K enterprise seats in 12 months."),
            ("Product Manager", "Atlassian", "2019 – 2022",
             "Owned Jira's automation feature. Shipped 40 workflows. 35% increase in daily active automations."),
        ],
        "education": [
            ("MBA", "INSEAD", "2019"),
            ("B.Eng. Computer Science", "American University of Beirut", "2016"),
        ],
    },

    # ── Marketing / Growth ────────────────────────────────────────────────────
    {
        "name": "Sophie Laurent",
        "email": "sophie.laurent@example.com",
        "phone": "+33 7 55 01 23 00",
        "location": "London, UK (Remote)",
        "summary": (
            "Growth Marketing Manager with 7 years scaling B2B SaaS products from pre-PMF to "
            "Series C. Specialises in demand generation, SEO/SEM, and product-led growth motions."
        ),
        "skills": "Performance Marketing, SEO/SEM, Google Ads, HubSpot, Salesforce, SQL, Looker, A/B Testing, Email Marketing, Content Strategy",
        "experience": [
            ("Head of Growth", "Lemon.io", "2022 – Present",
             "Led demand gen: grew qualified pipeline 3× YoY. Owned €2M paid media budget. CAC reduced 30%."),
            ("Growth Marketing Manager", "Contentful", "2019 – 2022",
             "Built SEO programme from 0: 500K organic monthly visits in 18 months. Managed PPC and marketing automation."),
        ],
        "education": [("M.Sc. Marketing & Communications", "Sciences Po Paris", "2017")],
    },

    # ── Business Operations ───────────────────────────────────────────────────
    {
        "name": "Jordan Lee",
        "email": "jordan.lee@example.com",
        "phone": "+1 415 555 1230",
        "location": "San Francisco, CA",
        "summary": (
            "Business Operations & Strategy Manager with 8 years at high-growth tech "
            "companies. Expert in operating cadences (OKRs, QBRs), cross-functional "
            "programme management, and data-driven decision making."
        ),
        "skills": "Business Operations, OKR Facilitation, Strategic Planning, SQL, Mode Analytics, Asana, Notion, Stakeholder Management, Headcount Planning, Board Reporting",
        "experience": [
            ("Senior Business Operations Manager", "Figma", "2022 – Present",
             "Designed and ran OKR programme across 800-person Engineering and Product org. Built executive dashboards in Mode."),
            ("Operations Manager", "Dropbox", "2019 – 2022",
             "Owned annual planning process for R&D. Partnered with Finance on $300M budget. Led operational efficiency initiative saving $12M."),
        ],
        "education": [
            ("MBA", "Stanford GSB", "2019"),
            ("B.Sc. Industrial Engineering", "UC Davis", "2016"),
        ],
    },
]

JD_TEXT = """Engineering Manager — Backend Infrastructure

About the Role
We are looking for an experienced Engineering Manager to lead our Backend Infrastructure team of 8–12 engineers. You will drive technical strategy, grow the team, and ensure reliable delivery of the systems that power our product.

Responsibilities
- Lead and grow a high-performing team of backend engineers
- Own the technical roadmap for backend infrastructure services
- Partner with Product and Design on feature delivery
- Drive engineering best practices: code review, testing, on-call culture
- Hire, mentor, and retain senior engineers
- Manage performance and career growth of direct reports
- Communicate progress, risks, and trade-offs to senior leadership

Required Skills & Experience
- 3+ years in an Engineering Manager or Tech Lead role
- 7+ years of software engineering experience
- Strong backend engineering background (Python, Go, or Java)
- Experience with distributed systems and cloud infrastructure (AWS preferred)
- Proven ability to hire and retain senior engineers
- Track record of shipping high-quality, high-scale software
- Strong written and verbal communication

Nice to Have
- Experience with ML infrastructure or data pipelines
- Exposure to event-driven architectures (Kafka, SQS)
- Startup and scale-up experience
- Experience with Kubernetes or container orchestration

We offer competitive compensation, remote-friendly working, and a strong engineering culture.
"""

def make_docx(candidate: dict, path: str) -> None:
    doc = Document()
    doc.add_heading(candidate["name"], 0)
    doc.add_paragraph(f"Email: {candidate['email']}  |  Phone: {candidate['phone']}  |  Location: {candidate['location']}")
    doc.add_heading("Summary", 1)
    doc.add_paragraph(candidate["summary"])
    doc.add_heading("Skills", 1)
    doc.add_paragraph(candidate["skills"])
    doc.add_heading("Experience", 1)
    for title, company, period, desc in candidate["experience"]:
        doc.add_paragraph(f"{title} — {company} ({period})", style="List Bullet")
        doc.add_paragraph(desc)
    doc.add_heading("Education", 1)
    for degree, institution, year in candidate["education"]:
        doc.add_paragraph(f"{degree}, {institution} ({year})", style="List Bullet")
    doc.save(path)
    print(f"  Created: {path}")


if __name__ == "__main__":
    print("Generating synthetic sample resumes...")
    for c in CANDIDATES:
        safe_name = c["name"].lower().replace(" ", "_")
        path = os.path.join(RESUMES_DIR, f"{safe_name}_resume.docx")
        make_docx(c, path)

    jd_path = os.path.join(JOBS_DIR, "engineering_manager_backend.txt")
    with open(jd_path, "w") as f:
        f.write(JD_TEXT)
    print(f"  Created: {jd_path}")

    print(f"\nDone. {len(CANDIDATES)} resumes + 1 JD in data/")
